import streamlit as st
import os
import tempfile
from datetime import datetime
import requests
import json
import base64
from docx import Document
from docx.shared import Inches  # docx.shared.Inches 문제 해결을 위한 임포트
import fitz  # PyMuPDF
import io
from PIL import Image
from openai import OpenAI
from dotenv import load_dotenv
import re

# ----- 유틸리티 함수 -----

def firecrawl_research(query, domains=None):
    """Firecrawl API를 사용한 웹 데이터 수집"""
    
    # Firecrawl API 호출
    try:
        # 최신 문서에 맞는 API 엔드포인트 (v1 버전 지정)
        url = "https://api.firecrawl.dev/v1/search"
        
        headers = {
            "Authorization": f"Bearer {firecrawl_api_key}",
            "Content-Type": "application/json"
        }
        
        # 개행 문자 제거 및 양쪽 공백 제거
        query = query.strip()
        
        # API 요청 페이로드 (v1 API 문서 기준 업데이트)
        payload = {
            "query": query  # ✅ 현재 Firecrawl v1에서 허용된 유일한 key
        }
        
        # 참고: v1 API에서는 'sites', 'num_results', 'lang', 'time_range' 필드가 지원되지 않음
        # domains 파라미터는 무시됨
        
        # API 호출 전 로그
        st.info(f"Firecrawl API 호출 중: {url}")
        st.info(f"API 키: {firecrawl_api_key[:5]}...{firecrawl_api_key[-5:]}")
        st.info(f"요청 페이로드: {payload}")
        
        response = requests.post(url, headers=headers, json=payload)
        
        if response.status_code == 200:
            result_data = response.json()
            
            # 디버깅용 응답 출력
            st.write("API 응답:", result_data.keys())
            
            # 결과 포맷팅
            formatted_results = {
                "text_content": [],
                "images": [],
                "sources": []
            }
            
            # 텍스트 컨텐츠 추출 (Firecrawl v1 API 응답 형식에 맞게 조정)
            results = result_data.get("data", [])  # 'data' 필드 확인
            
            if not results:
                # 다른 가능한 필드들도 확인
                results = result_data.get("results", [])
                
                if not results and "organic" in result_data:
                    # v1 API에서는 'organic' 필드 아래에 결과가 있을 수 있음
                    results = result_data.get("organic", [])
                
            for item in results:
                # 스니펫 추출
                if "snippet" in item:
                    formatted_results["text_content"].append(item["snippet"])
                    
                # 제목 추출 (v1 API에서는 다른 필드명일 수 있음)
                if "title" in item:
                    formatted_results["text_content"].append(f"제목: {item['title']}")
                    
                # URL 추출    
                if "link" in item:  # v1 API에서는 'url' 대신 'link'일 수 있음
                    formatted_results["sources"].append(item["link"])
                elif "url" in item:
                    formatted_results["sources"].append(item["url"])
                    
                # 이미지 추출    
                if "image" in item:  # v1 API에서는 'image_url' 대신 'image'일 수 있음
                    title = item.get("title", "이미지")
                    formatted_results["images"].append((item["image"], title))
                elif "image_url" in item and "title" in item:
                    formatted_results["images"].append((item["image_url"], item["title"]))
            
            # 응답 메시지
            st.success(f"Firecrawl API 검색 결과: {len(results)}건 조회됨")
            
            # 응답이 비어 있으면 예시 데이터 사용
            if not formatted_results["text_content"]:
                st.warning("검색 결과가 없어 예시 데이터를 사용합니다.")
                return get_example_data()
                
            return formatted_results
        else:
            # API 오류시 상세 정보 표시 및 예시 데이터 반환
            error_msg = f"API 오류 (예시 데이터 사용): {response.status_code}, {response.text}"
            st.warning(error_msg)
            print(error_msg)
            
            # API 키 검증 문제인 경우
            if response.status_code == 401:
                st.error("API 키가 유효하지 않습니다. 환경 변수 FIRECRAWL_API_KEY를 확인하세요.")
                
            return get_example_data()
    except Exception as e:
        # 예외 발생시 상세 정보 표시
        error_msg = f"API 연결 오류 (예시 데이터 사용): {str(e)}"
        st.warning(error_msg)
        print(error_msg)
        return get_example_data()

def get_example_data():
    """데모용 예시 데이터"""
    st.info("예시 데이터를 사용합니다.")
    return {
        "text_content": [
            "생성형 AI 시장은 2025년까지 연간 38% 성장할 것으로 예상된다.",
            "기업들은 AI를 활용한 업무 자동화에 투자를 확대하고 있다.",
            "ChatGPT와 같은 대화형 AI는 고객 서비스 분야에서 혁신을 일으키고 있다.",
            "Google DeepMind의 최신 연구에 따르면 AI 모델의 성능이 매년 2배씩 향상되고 있다."
        ],
        "images": [
            ("https://mblogthumb-phinf.pstatic.net/MjAyMzA0MThfMTA0/MDAxNjgxNzg2NDk1MzM4.Cz-AbXLhvkdrYGEQXS_D7P3y3rQooUW-pRXCIhKJEnAg.a5B2vWMe8YdBOZENDBmJ-VlE-vGnI2qJKBYx6YfhZ2Ig.PNG.esaracen/IMG_2.png", "AI 시장 성장 예측 그래프"),
            ("https://blog.kakaocdn.net/dn/bYoLUJ/btr8nQs0KRy/UXgAC0AWypxQJLKlkSrDDK/img.png", "ChatGPT 인터페이스")
        ],
        "sources": [
            "https://techcrunch.com/2023/05/15/ai-market-growth",
            "https://www.mckinsey.com/ai-adoption-survey-2023",
            "https://research.google/blog/ai-trends-2024"
        ]
    }

def extract_from_pdf(file_path):
    """PDF 파일에서 텍스트와 이미지 추출"""
    result = {
        "text_content": [],
        "images": []
    }
    
    try:
        pdf_document = fitz.open(file_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # 텍스트 추출
            text = page.get_text()
            if text.strip():
                result["text_content"].append(text)
            
            # 이미지 추출
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # 이미지 처리 및 저장
                image = Image.open(io.BytesIO(image_bytes))
                img_buffer = io.BytesIO()
                image.save(img_buffer, format="PNG")
                img_data = base64.b64encode(img_buffer.getvalue()).decode("utf-8")
                
                result["images"].append((f"data:image/png;base64,{img_data}", f"이미지 {page_num+1}-{img_index+1}"))
        
        pdf_document.close()
    except Exception as e:
        error_msg = f"PDF 추출 오류: {str(e)}"
        st.error(error_msg)
        print(error_msg)
    
    return result

def extract_from_docx(file_path):
    """Word 문서에서 텍스트와 이미지 추출"""
    result = {
        "text_content": [],
        "images": []
    }
    
    try:
        doc = Document(file_path)
        
        # 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip():
                result["text_content"].append(para.text)
        
        # 이미지 추출 (추가 구현 필요)
        # python-docx에서는 이미지 추출이 직접적으로 지원되지 않음
        # 필요 시 추가 라이브러리 사용해 구현
        
    except Exception as e:
        error_msg = f"DOCX 추출 오류: {str(e)}"
        st.error(error_msg)
        print(error_msg)
    
    return result

def generate_report(user_query, collected_data, style="기사형", include_title=True, 
                   include_lead=True, include_body=True, include_sources=True, 
                   report_length=2, temperature=0.3):
    """GPT를 이용한 보고서 생성"""
    
    # 프롬프트 구성
    prompt = "다음 정보를 기반으로 "
    
    if style == "기사형":
        prompt += "경제/기술 기사처럼 작성해주세요. 문장은 간결하고, 정보 중심적으로 구성하며, 독자가 빠르게 요점을 이해할 수 있도록 구성해주세요."
    else:
        prompt += f"다음 스타일로 작성해주세요: {style}"
    
    # 포함할 항목 지정
    report_sections = []
    if include_title:
        report_sections.append("제목 (강력한 메시지를 담은 헤드라인)")
    if include_lead:
        report_sections.append("리드 문단 (핵심 요약 또는 임팩트 있는 문장)")
    if include_body:
        report_sections.append("주요 본문 (단락으로 구분된 상세 내용)")
    if include_sources:
        report_sections.append("참고 링크 목록")
    
    prompt += f"\n\n포함할 항목: {', '.join(report_sections)}"
    
    # 분량 설정
    length_map = {
        1: "매우 간결하게 (300단어 이내)",
        2: "간결하게 (500단어 이내)",
        3: "보통 분량 (800단어 이내)",
        4: "상세하게 (1200단어 이내)",
        5: "매우 상세하게 (2000단어 이내)"
    }
    
    prompt += f"\n\n분량: {length_map[report_length]}"
    
    # 검색 질의 및 수집 데이터 추가
    prompt += f"\n\n사용자 질의: {user_query}\n\n"
    prompt += "수집된 정보:\n"
    
    for i, text in enumerate(collected_data["text_content"]):
        prompt += f"{i+1}. {text}\n"
    
    if collected_data["sources"]:
        prompt += "\n참고 링크:\n"
        for source in collected_data["sources"]:
            prompt += f"- {source}\n"
    
    # 링크 목록 형식을 명확히 지정
    if include_sources and collected_data["sources"]:
        prompt += "\n마지막에 반드시 '참고 링크:' 제목 아래에 수집된 출처 링크를 목록 형태로 포함시켜주세요. 각 링크는 새 줄에 표시하고 하이퍼링크 형식으로 작성하세요."
    
    # OpenAI API 호출 (1.0.0 이상 버전 방식)
    try:
        client = OpenAI(api_key=openai_api_key)
        
        response = client.chat.completions.create(
            model="gpt-4",  # 또는 다른 모델
            messages=[
                {"role": "system", "content": "당신은 전문 기자이자 리서치 애널리스트입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=temperature,
            max_tokens=3000
        )
        return response.choices[0].message.content
    except Exception as e:
        error_msg = f"GPT API 오류: {str(e)}"
        st.error(error_msg)
        print(error_msg)
        return f"보고서 생성 중 오류가 발생했습니다: {str(e)}"

def format_report_with_links(report_content, sources):
    """보고서 내용에 참고 링크가 없는 경우 추가"""
    if '참고 링크:' not in report_content and sources:
        report_content += "\n\n## 참고 링크:\n"
        for source in sources:
            report_content += f"- [{source}]({source})\n"
    
    return report_content

def create_docx_report(report_content, images=None):
    """DOCX 형식 보고서 생성"""
    doc = Document()
    
    # 마크다운을 기본 텍스트로 변환 (간단한 처리)
    lines = report_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        if line.startswith('# '):
            # 큰 제목
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            # 중간 제목
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            # 작은 제목
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            # 목록
            doc.add_paragraph(line[2:], style='ListBullet')
        elif line.strip() == '':
            # 빈 줄
            if current_paragraph:
                current_paragraph = None
        else:
            # 일반 문단
            if not current_paragraph:
                current_paragraph = doc.add_paragraph()
            current_paragraph.add_run(line)
    
    # 이미지 추가 (이미지가 있는 경우)
    if images:
        doc.add_heading('관련 이미지', level=1)
        for img_url, caption in images:
            try:
                if img_url.startswith('data:image'):
                    # Base64 인코딩된 이미지
                    img_data = img_url.split(',')[1]
                    binary_img = base64.b64decode(img_data)
                    img_stream = io.BytesIO(binary_img)
                    doc.add_picture(img_stream, width=Inches(5))  # docx.shared.Inches → Inches로 수정
                else:
                    # URL 이미지 (실제로는 다운로드 필요)
                    response = requests.get(img_url)
                    img_stream = io.BytesIO(response.content)
                    doc.add_picture(img_stream, width=Inches(5))  # docx.shared.Inches → Inches로 수정
                
                # 캡션 추가 (스타일 이름 사용)
                doc.add_paragraph(caption, style='Caption')
            except Exception as e:
                error_msg = f"이미지 추가 오류: {str(e)}"
                print(error_msg)
    
    # 메모리 스트림에 저장
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream.getvalue()

def create_pdf_report(report_content, images=None):
    """PDF 형식 보고서 생성 (간단한 구현)"""
    # 실제 구현은 ReportLab 또는 WeasyPrint 사용 필요
    # 지금은 예시 데이터 반환
    pdf_data = io.BytesIO()
    pdf_data.write(b"PDF example data")
    pdf_data.seek(0)
    return pdf_data.getvalue()

# 환경 변수 로드
load_dotenv()

# API 키 설정
openai_api_key = os.getenv("OPENAI_API_KEY")
firecrawl_api_key = os.getenv("FIRECRAWL_API_KEY")

# 페이지 기본 설정
st.set_page_config(
    page_title="심층 웹 리서치 보고서 생성기",
    page_icon="🔍",
    layout="wide",
)

# 세션 상태 초기화
if "generated_report" not in st.session_state:
    st.session_state.generated_report = None
if "report_images" not in st.session_state:
    st.session_state.report_images = []
if "report_sources" not in st.session_state:
    st.session_state.report_sources = []
if "progress" not in st.session_state:
    st.session_state.progress = 0

# 앱 제목 및 설명
st.title("🔍 심층 웹 리서치 자동 보고서 생성기")
st.markdown("웹 데이터 또는 문서를 분석하여 기사형 보고서를 자동으로 생성합니다.")

# 사이드바 설정
with st.sidebar:
    st.header("설정")
    
    # 문체 선택
    st.subheader("문체 선택")
    style_option = st.radio(
        "보고서 스타일:",
        ["기사형 (기본)", "직접 입력"],
        index=0
    )
    
    custom_style = ""
    if style_option == "직접 입력":
        custom_style = st.text_area(
            "스타일 지정을 위한 예시 텍스트 또는 지시사항을 입력하세요:",
            height=100
        )
    
    # 보고서 항목 선택
    st.subheader("보고서 항목")
    include_title = st.checkbox("제목", value=True)
    include_lead = st.checkbox("리드 문단", value=True)
    include_body = st.checkbox("주요 본문", value=True)
    include_images = st.checkbox("이미지 포함", value=True)
    include_sources = st.checkbox("참고 링크", value=True)
    
    # 분량 설정
    st.subheader("분량 설정")
    report_length = st.slider("보고서 길이", min_value=1, max_value=5, value=2, 
                              help="1: 매우 짧게, 5: 매우 자세하게")
    
    # 추가 옵션
    st.subheader("고급 설정")
    temperature = st.slider("창의성 수준", min_value=0.0, max_value=1.0, value=0.3, step=0.1,
                            help="낮을수록 일관된 결과, 높을수록 창의적인 결과")

# 탭 설정
tab1, tab2 = st.tabs(["리서치 입력", "결과 보고서"])

# 리서치 입력 탭
with tab1:
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.subheader("질문 또는 주제")
        user_query = st.text_area("연구하고 싶은 질문이나 주제를 입력하세요", height=100)
        
        st.subheader("웹 데이터 설정 (Firecrawl)")
        reference_domains = st.text_input(
            "참조할 도메인 (쉼표로 구분, 빈칸이면 모든 사이트 검색)",
            placeholder="예: techcrunch.com, korea.kr, naver.com"
        )
    
    with col2:
        st.subheader("문서 업로드 (선택사항)")
        uploaded_file = st.file_uploader("PDF 또는 Word 문서 업로드", type=["pdf", "docx"])
        
        if uploaded_file:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            st.success(f"파일 업로드 완료: {uploaded_file.name}")
    
    # 실행 버튼
    if st.button("보고서 생성하기", type="primary"):
        if not user_query and not uploaded_file:
            st.error("질문을 입력하거나 문서를 업로드해주세요.")
        else:
            # 진행 상황 표시
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # 1. 데이터 수집 단계
            status_text.text("1/3 단계: 데이터 수집 중...")
            
            collected_data = {
                "text_content": [],
                "images": [],
                "sources": []
            }
            
            # 1-A: Firecrawl API를 통한 웹 데이터 수집 (질문이 있는 경우)
            if user_query:
                try:
                    status_text.text("웹 데이터 수집 중...")
                    
                    # Firecrawl API 호출 (개행 문자 및 공백 제거)
                    firecrawl_data = firecrawl_research(user_query.strip(), reference_domains)
                    
                    collected_data["text_content"].extend(firecrawl_data.get("text_content", []))
                    collected_data["images"].extend(firecrawl_data.get("images", []))
                    collected_data["sources"].extend(firecrawl_data.get("sources", []))
                    
                    progress_bar.progress(0.3)
                except Exception as e:
                    st.error(f"웹 데이터 수집 중 오류 발생: {str(e)}")
            
            # 1-B: 업로드된 문서 처리 (파일이 있는 경우)
            if uploaded_file:
                try:
                    status_text.text("문서 분석 중...")
                    
                    # 파일 형식에 따라 처리
                    if uploaded_file.name.lower().endswith('.pdf'):
                        pdf_data = extract_from_pdf(tmp_file_path)
                        collected_data["text_content"].extend(pdf_data.get("text_content", []))
                        collected_data["images"].extend(pdf_data.get("images", []))
                    
                    elif uploaded_file.name.lower().endswith('.docx'):
                        docx_data = extract_from_docx(tmp_file_path)
                        collected_data["text_content"].extend(docx_data.get("text_content", []))
                        collected_data["images"].extend(docx_data.get("images", []))
                    
                    # 파일명 출처로 추가
                    collected_data["sources"].append(f"업로드 문서: {uploaded_file.name}")
                    
                    progress_bar.progress(0.5)
                except Exception as e:
                    st.error(f"문서 분석 중 오류 발생: {str(e)}")
            
            # 2. 보고서 생성 단계
            if collected_data["text_content"]:
                try:
                    status_text.text("2/3 단계: 보고서 생성 중...")
                    
                    # GPT를 이용한 보고서 생성
                    report_content = generate_report(
                        user_query=user_query,
                        collected_data=collected_data,
                        style=custom_style if style_option == "직접 입력" else "기사형",
                        include_title=include_title,
                        include_lead=include_lead,
                        include_body=include_body,
                        include_sources=include_sources,
                        report_length=report_length,
                        temperature=temperature
                    )
                    
                    # 링크 목록 형식 확인 및 추가
                    report_content = format_report_with_links(report_content, collected_data["sources"])
                    
                    # 세션 상태에 보고서와 소스 저장
                    st.session_state.generated_report = report_content
                    st.session_state.report_sources = collected_data["sources"]
                    
                    # 이미지 처리 (include_images가 True인 경우)
                    if include_images and collected_data["images"]:
                        st.session_state.report_images = collected_data["images"][:3]  # 최대 3개만 사용
                    
                    progress_bar.progress(0.8)
                except Exception as e:
                    st.error(f"보고서 생성 중 오류 발생: {str(e)}")
            
            # 3. 결과 완료
            status_text.text("3/3 단계: 결과 정리 중...")
            progress_bar.progress(1.0)
            status_text.text("보고서 생성 완료! '결과 보고서' 탭을 확인하세요.")
            
            # 결과 탭으로 자동 전환
            st.query_params.active_tab = "result"

# 결과 보고서 탭
with tab2:
    if st.session_state.generated_report:
        st.markdown("## 생성된 보고서")
        
        # 보고서 내용 표시
        st.markdown(st.session_state.generated_report, unsafe_allow_html=True)
        
        # 이미지 표시 (있는 경우)
        if st.session_state.report_images and len(st.session_state.report_images) > 0:
            st.subheader("관련 이미지")
            image_cols = st.columns(min(3, len(st.session_state.report_images)))
            
            for i, (img_url, caption) in enumerate(st.session_state.report_images):
                if i < len(image_cols):
                    with image_cols[i]:
                        try:
                            # 이미지 URL 디버깅
                            st.write(f"이미지 로드 중: {img_url[:50]}...")
                            st.image(img_url, caption=caption, use_column_width=True)
                        except Exception as e:
                            st.error(f"이미지 로드 실패: {str(e)}")
        else:
            st.info("관련 이미지가 없습니다.")
        
        # 참고 링크 표시
        if st.session_state.report_sources and len(st.session_state.report_sources) > 0:
            st.subheader("참고 링크 목록")
            for source in st.session_state.report_sources:
                st.markdown(f"- [{source}]({source})")
        
        # 다운로드 버튼
        col1, col2 = st.columns(2)
        
        with col1:
            # DOCX 형식 다운로드
            docx_file = create_docx_report(
                st.session_state.generated_report, 
                st.session_state.report_images
            )
            
            st.download_button(
                label="DOCX 형식으로 다운로드",
                data=docx_file,
                file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col2:
            # PDF 형식 다운로드 (추후 구현 예정)
            st.download_button(
                label="PDF 형식으로 다운로드",
                data=create_pdf_report(st.session_state.generated_report, st.session_state.report_images),
                file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
    else:
        st.info("보고서가 아직 생성되지 않았습니다. '리서치 입력' 탭에서 보고서를 생성해주세요.") 